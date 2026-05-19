"""Generate BandMS Admin Panel functional description as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def table_2col(rows, header=None):
    cols = 2
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    t.style = 'Table Grid'
    if header:
        for i, h in enumerate(header):
            c = t.rows[0].cells[i]
            c.text = h
            for run in c.paragraphs[0].runs:
                run.bold = True
        for i, (a, b) in enumerate(rows):
            row = t.rows[i + 1]
            row.cells[0].text = a
            row.cells[1].text = b
    else:
        for i, (a, b) in enumerate(rows):
            t.rows[i].cells[0].text = a
            t.rows[i].cells[1].text = b
    return t

def table_ncol(data, headers):
    t = doc.add_table(rows=len(data) + 1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            t.rows[ri + 1].cells[ci].text = val
    return t

def spacer():
    doc.add_paragraph()

# ── Title page ─────────────────────────────────────────────────────────────────

title = doc.add_heading('BandMS Admin Panel', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Functional Description & Test Scenario Reference')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ── TOC placeholder ────────────────────────────────────────────────────────────

h1('Contents')
sections = [
    '1. Overview & Access Control',
    '2. Admin Dashboard',
    '3. Band Profile (6 tabs)',
    '4. Band Members',
    '5. Bands',
    '6. Venues',
    '7. Concerts',
    '8. Posts',
    '9. Tags',
    '10. Photos & Albums',
    '11. Releases',
    '12. Tours',
    '13. Press Releases',
    '14. Music Videos',
    '15. Instruments',
    '16. Authors / Contacts',
    '17. Pitch Generator',
    '18. Band Calendar',
    '19. Tech Rider (Admin only)',
    '20. Newsletter (Admin only)',
    '21. Users (Admin only)',
    '22. Setlists (Admin only)',
    '23. My Profile (Member)',
    '24. My Setups (Member)',
    '25. Shared UI Patterns',
    '26. Common Workflows',
]
for s in sections:
    bullet(s)

doc.add_page_break()

# ── 1. Overview ────────────────────────────────────────────────────────────────

h1('1. Overview & Access Control')
body(
    'The BandMS admin panel is a Vue 3 single-page application served at /admin. '
    'Every route requires authentication via a Passport Bearer token. '
    'Three roles exist: Admin, Member, and Publisher. '
    'Most pages are accessible to all authenticated users; a few are restricted to Admin or Member.'
)
spacer()
h2('Route & Role Matrix')
table_ncol([
    ['/admin', 'AdminDashboard', 'All users', 'Stats, EPK versioning, career level, press widget'],
    ['/admin/band-profile', 'BandProfileAdminView', 'All users', '6-tab profile editor'],
    ['/admin/band-members', 'BandMembersAdminView', 'All users', 'Dual-pane member list + detail'],
    ['/admin/bands', 'BandsAdminView', 'All users', 'Other bands (concert lineups)'],
    ['/admin/venues', 'VenuesAdminView', 'All users', 'Venue CRUD'],
    ['/admin/concerts', 'ConcertsAdminView', 'All users', 'Concert CRUD + poster upload'],
    ['/admin/posts', 'PostsAdminView', 'All users', 'Blog posts CRUD'],
    ['/admin/tags', 'TagsAdminView', 'All users', 'Global tag CRUD'],
    ['/admin/photos', 'PhotosAdminView', 'All users', 'Album + photo manager'],
    ['/admin/releases', 'ReleasesAdminView', 'All users', 'Discography CRUD'],
    ['/admin/tours', 'ToursAdminView', 'All users', 'Tour CRUD'],
    ['/admin/press-releases', 'PressReleasesAdminView', 'All users', 'Press articles CRUD'],
    ['/admin/music-videos', 'MusicVideosAdminView', 'All users', 'YouTube video manager'],
    ['/admin/instruments', 'InstrumentsAdminView', 'All users', 'Instrument catalog CRUD'],
    ['/admin/authors', 'AuthorsAdminView', 'All users', 'Journalist/contact CRUD'],
    ['/admin/pitch', 'PitchGeneratorView', 'All users', '6 pitch template generator'],
    ['/admin/band-calendar', 'BandCalendarView', 'All users', 'FullCalendar from member URLs'],
    ['/admin/tech-rider', 'TechRiderAdminView', 'Admin only', 'Multi-section rider builder'],
    ['/admin/newsletter', 'NewsletterAdminView', 'Admin only', 'Subscriber list + CSV export'],
    ['/admin/users', 'UsersAdminView', 'Admin only', 'User account management'],
    ['/admin/setlists', 'SetlistsAdminView', 'Admin only', 'Setlist builder + song library'],
    ['/admin/my-profile', 'MyProfileView', 'Member (own)', 'Self-edit band member profile'],
    ['/admin/my-setups', 'MySetupsView', 'Member (own)', 'Self-manage stage setups'],
], ['Path', 'Component', 'Access', 'Purpose'])

doc.add_page_break()

# ── 2. Dashboard ───────────────────────────────────────────────────────────────

h1('2. Admin Dashboard')
body('The dashboard is the landing page after login. It provides a quick overview of the band\'s data and surfaces actionable items.')
spacer()

h2('Stats Grid')
body('Seven cards arranged in a grid, each linking to the corresponding admin section:')
for item in ['Bands', 'Releases', 'Tours', 'Venues', 'Concerts', 'Tags', 'Posts']:
    bullet(item)
spacer()

h2('EPK Versions Widget')
body('Shows the current EPK publication status.')
bullet('Displays "Live" version number and date, or "No published version" if none exists.')
bullet('If a pending version exists: version number, creation date, and optional reason text.')
bullet('Actions: Publish (green button), Discard (red button).')
bullet('Link to "Create snapshot" navigates to Band Profile → EPK Settings tab.')
spacer()

h2('Career Level Widget')
body('Visual checklist of career milestones mapped to a 1–4 level system.')
bullet('Level selector: 4 clickable buttons (Garage / Local / Pro / Custom).')
bullet('Checklist items: concerts count, releases, press articles, posts, music videos, members, tech rider active, EPK published.')
bullet('Clicking a level button saves the new level immediately via PUT /api/band-profile.')
spacer()

h2('Press Coverage Enhancement Widget')
body('Shows the top 5 press releases with the lowest metadata score.')
bullet('Score calculated from: og_title, og_image, featured flag, tags (max 3), linked concerts, releases, albums, tours, posts, authors (max 13 points total).')
bullet('Progress ring shows overall % of press releases considered "enhanced".')
bullet('Clicking a bar navigates to that press release in /admin/press-releases.')
spacer()

h2('API Calls')
for api in [
    'GET /api/band-profile',
    'GET /api/epk-versions  (+ publish/discard mutations)',
    'GET /api/bands, /api/venues, /api/concerts, /api/tags, /api/releases, /api/tours  (stats cards)',
    'GET /api/press-releases, /api/posts, /api/music-videos, /api/tech-riders  (widgets)',
]:
    bullet(api)

doc.add_page_break()

# ── 3. Band Profile ────────────────────────────────────────────────────────────

h1('3. Band Profile')
body('Six-tab editor at /admin/band-profile. A single Save button at the bottom submits the currently visible tab\'s data via PUT /api/band-profile (partial update — unset fields are not cleared).')
spacer()

tabs = [
    ('Tab 1: Bio', [
        'Band name (required, max 255 chars).',
        'Bio section with 4 sub-tabs: One-liner (max 280 chars with live counter), Short (plain textarea), Long (rich editor), Full (rich editor).',
        'Character warning shown at 240 chars for the one-liner field.',
        'Tab header shows red indicator dot if the active bio sub-tab has a validation error.',
    ]),
    ('Tab 2: Career', [
        'Band level selector — 4 cards (Garage / Local / Pro / Custom) with emoji icons; clicking saves immediately.',
        'Formation year (number input).',
        'Hometown / Base (text).',
        'Genres (comma-separated, displayed as badges on EPK).',
        'Comparable artists (comma-separated).',
        'Artistic statement (textarea).',
    ]),
    ('Tab 3: Social Links', [
        'List of existing social links with platform colour dot, label, URL, and Edit / Delete buttons.',
        'Supported platforms: Instagram, Facebook, YouTube, Spotify, TikTok, Twitter/X, Bandcamp, SoundCloud, Website.',
        '"Add link" button reveals an inline form (platform dropdown + URL input + Save/Cancel).',
        'Edit button re-opens the inline form pre-filled for that link.',
        'Delete button triggers inline confirmation before removing.',
        'API: GET /api/band-profile/social-links, POST, PUT /{id}, DELETE /{id}.',
    ]),
    ('Tab 4: Contacts', [
        'Booking email (validated as email).',
        'Press / PR email.',
        'Tech contact email.',
        'Tech contact phone.',
        'Sound engineer description (textarea, appears on tech rider cover page).',
    ]),
    ('Tab 5: Stats', [
        'Facebook likes — text input + "Sync from Facebook" button (POST /api/band-profile/sync-facebook); shows last sync timestamp.',
        'Spotify monthly listeners (manual number input).',
        'Instagram followers (manual).',
        'TikTok followers (manual).',
        'YouTube subscribers (manual).',
        'Facebook followers (manual).',
    ]),
    ('Tab 6: EPK Settings', [
        'Featured release — dropdown populated from /api/releases.',
        'EPK Snapshot — "Create EPK snapshot" button opens modal with optional reason/notes field; submits POST /api/epk-versions.',
        'Tech rider PDF — upload button (max 10 MB), existing file shown with filename + Remove button.',
        'Stage plot image — upload button (max 4 MB), thumbnail preview if exists, Remove button.',
        'API: POST /api/band-profile/upload-tech-rider, DELETE /api/band-profile/tech-rider, POST /api/band-profile/upload-stage-plot, DELETE /api/band-profile/stage-plot.',
    ]),
]
for title_text, items in tabs:
    h2(title_text)
    for item in items:
        bullet(item)
    spacer()

doc.add_page_break()

# ── 4. Band Members ────────────────────────────────────────────────────────────

h1('4. Band Members')
body('Dual-pane layout: fixed 240 px sidebar on the left, detail pane on the right.')
spacer()

h2('Left Sidebar — Member List')
bullet('Header: title + "+" button to add a new member.')
bullet('Two sections: "Current lineup" (is_current: true) and "Former" (is_current: false).')
bullet('Each card: avatar (photo or initials), full name, role, instrument list.')
bullet('Click a card to load that member in the detail pane.')
bullet('Drag-and-drop reorder within each section; on drop, sends PUT /api/band-profile/members/reorder with new ID array.')
spacer()

h2('Right Detail Pane')
body('Empty state with 🎸 icon and "Add first member" prompt when nothing is selected.')
body('When a member is selected:')
bullet('Top bar: avatar, name, role label, "Remove" button (opens ConfirmDialog).')
bullet('Three tabs below the top bar:')
spacer()

h3('Tab 1 — Profile')
for item in [
    'First name, last name (required).',
    'Photo upload / removal.',
    'Bio (textarea).',
    'Role (text, e.g. "Lead Vocals").',
    'Instruments — multi-select from /api/instruments.',
    'Social links (same inline form as Band Profile Social tab).',
    'Joined date / Quit date (date pickers; quit must be ≥ joined).',
    'Is current member (checkbox).',
    'Nickname (text).',
]:
    bullet(item)
spacer()

h3('Tab 2 — Stage Setups')
for item in [
    'List of setups for this member with per-setup sub-tabs.',
    '"New setup" button creates a blank setup.',
    'Each setup: name, signal chain type, inputs table, monitor mix, backline items, power requirements.',
    'Delete setup button (with confirm).',
]:
    bullet(item)
spacer()

h3('Tab 3 — Default Gear')
for item in [
    'Default backline preferences.',
    'Default power requirements.',
    'Default monitor preferences.',
    'Saved independently via PUT /api/band-members/{id} with default_gear payload.',
]:
    bullet(item)
spacer()

h2('Add Member Modal')
bullet('"+" button opens AdminModal with blank Profile form.')
bullet('On save: POST /api/band-profile/members; modal closes; new member selected in sidebar.')
spacer()

h2('API Calls')
for api in [
    'GET /api/band-profile/members',
    'POST /api/band-profile/members',
    'PUT /api/band-profile/members/{id}',
    'DELETE /api/band-profile/members/{id}',
    'POST /api/band-profile/members/{id}/photo  (file upload)',
    'PUT /api/band-profile/members/reorder  (array of IDs)',
    'GET /api/instruments  (for multi-select)',
]:
    bullet(api)

doc.add_page_break()

# ── Standard CRUD section ──────────────────────────────────────────────────────

h1('5–9. Standard CRUD Views (Bands, Venues, Concerts, Posts, Tags)')
body(
    'These five views share the same layout pattern: sortable table with search, '
    'create/edit modal, and delete confirmation dialog.'
)
spacer()

h2('Shared Layout Pattern')
bullet('Header row: page title + "Add [resource]" button.')
bullet('TableToolbar: real-time search input (client-side filter), optional filter dropdown, "Showing X of Y" counter.')
bullet('Sortable table: clicking a column header toggles ascending / descending sort. Active sort shown with arrow icon.')
bullet('Pagination: Prev / Next buttons, current page indicator. Shown only when total > 15 rows.')
bullet('Empty state when no rows match the active search.')
spacer()

h2('Create / Edit Modal')
bullet('Opens AdminModal overlay with the resource-specific form pre-filled (edit) or blank (create).')
bullet('Field-level validation errors from the API are displayed below each input.')
bullet('Submit button shows loading spinner during the request; disabled until idle.')
bullet('Modal closes automatically on success; stays open with errors on failure.')
spacer()

h2('Delete Confirmation')
bullet('ConfirmDialog overlay: "Are you sure?" message, Confirm (red) and Cancel buttons.')
bullet('Row is removed from the table optimistically; error toast shown if the API call fails.')
spacer()

h2('Per-View Specifics')
table_ncol([
    ['Bands', '/admin/bands', 'Name, Genre, Actions', '"Message" quick-action opens Pitch Generator with band pre-filled as recipient.'],
    ['Venues', '/admin/venues', 'Name, City, Street, Actions', 'Street address split into street + number + city + postcode fields in the form.'],
    ['Concerts', '/admin/concerts', 'Date, Doors/Start time, Venue, Lineup, Actions', 'Upcoming / Past filter. Poster image upload in the edit form. Multi-select band lineup.'],
    ['Posts', '/admin/posts', 'Title, Tags, Published, Actions', 'Draft / Published filter. Rich editor for body. Tag multi-select. Related content panel.'],
    ['Tags', '/admin/tags', 'Name, Slug, Actions', 'Slug auto-generated from name but editable. Used across venues, posts, photos, concerts, etc.'],
], ['View', 'Path', 'Columns', 'Notes'])

doc.add_page_break()

# ── 10. Photos ─────────────────────────────────────────────────────────────────

h1('10. Photos & Albums')
body('Album-centric photo manager at /admin/photos.')
spacer()

h2('Batch Upload Modal')
bullet('"Batch upload" button opens modal with multi-file input.')
bullet('Metadata form: title, description, venue (dropdown), concert (dropdown), tags (multi-select), taken date.')
bullet('Upload progress bar shown during submit.')
bullet('Creates an album + uploads all selected photos in one API call.')
spacer()

h2('Album List')
bullet('Each album card: title, description, venue, concert, taken date, photo count.')
bullet('Edit button opens the Edit Album modal.')
bullet('Delete button removes the album and all its photos (with confirm).')
spacer()

h2('Edit Album Modal')
bullet('Metadata fields: title, description, venue, concert, taken date, published date, tags.')
bullet('Photo grid: thumbnails with drag-reorder, per-photo EPK-featured toggle (star icon), per-photo caption input, per-photo delete button.')
bullet('Save saves album metadata + new order + updated captions in one request.')
spacer()

h2('API Calls')
for api in [
    'POST /api/albums/batch-create  (multipart with files + metadata)',
    'PUT /api/albums/{id}',
    'DELETE /api/albums/{id}',
    'POST /api/albums/{id}/reorder-photos',
    'DELETE /api/albums/{id}/photos/{photoId}',
    'PUT /api/photos/{id}/toggle-epk-featured',
]:
    bullet(api)

doc.add_page_break()

# ── 11. Releases ───────────────────────────────────────────────────────────────

h1('11. Releases')
body('Discography manager at /admin/releases.')
spacer()

h2('Table')
bullet('Type filter dropdown: All / LP / EP / Single / Compilation.')
bullet('Columns: Type (colour-coded badge), Title, Release date, Photos count, Actions.')
bullet('Edit and Delete actions per row.')
spacer()

h2('Create / Edit Form')
bullet('Title (required).')
bullet('Artist (text, defaults to band name).')
bullet('Type dropdown: LP / EP / Single / Compilation.')
bullet('Release date (date picker).')
bullet('Description (rich editor).')
bullet('Genres (comma-separated).')
bullet('Credits (comma-separated).')
bullet('Cover image upload / replace / remove (POST /api/releases/{id}/upload-cover).')
bullet('Associated photos: add from photo library, drag-reorder, remove.')
bullet('Streaming links list (platform + URL, same inline form as social links).')
spacer()

h2('API Calls')
for api in [
    'GET /api/releases',
    'POST /api/releases',
    'PUT /api/releases/{id}',
    'DELETE /api/releases/{id}',
    'POST /api/releases/{id}/upload-cover',
    'DELETE /api/releases/{id}/cover',
    'POST /api/releases/{id}/photos  (add from library)',
    'DELETE /api/releases/{id}/photos/{photoId}',
    'POST /api/releases/{id}/reorder-photos',
]:
    bullet(api)

doc.add_page_break()

# ── 12–14 Tours, Press Releases, Music Videos ──────────────────────────────────

h1('12. Tours')
body('Tour list at /admin/tours. Columns: Poster, Name, Start–End dates, Concert count.')
bullet('Create / Edit modal: name, poster image upload, start/end dates, description.')
bullet('Associated concerts shown as read-only list in edit modal.')
spacer()

h1('13. Press Releases')
body('Press article manager at /admin/press-releases.')
spacer()
h2('Table')
bullet('Columns: OG image thumbnail (or 📰 placeholder), Article title + site name, Tags, Published date, EPK badge if featured.')
spacer()
h2('Create / Edit Modal')
bullet('URL field (required) + "Fetch metadata" button — auto-populates og_title, og_image, og_description.')
bullet('Published date picker.')
bullet('Featured flag toggle (appears in EPK).')
bullet('Tags multi-select.')
bullet('Relations panel: link to concerts, albums, releases, tours, posts, authors.')
spacer()

h1('14. Music Videos')
body('YouTube video manager at /admin/music-videos.')
spacer()
bullet('Status filter: All / Published / Draft.')
bullet('"Sync view counts" button — POST /api/music-videos/sync-views — fetches latest YouTube view counts for all saved videos; toast shows "Synced X videos · Y total views".')
bullet('Table columns: Title, URL, View count, Synced at, Published date, Actions.')
bullet('Create / Edit modal: Title, YouTube URL, Published date / Draft toggle, Sort order, "Fetch preview" button (scrapes OG meta from the URL).')

doc.add_page_break()

# ── 15–16 Instruments, Authors ─────────────────────────────────────────────────

h1('15. Instruments')
body('Instrument catalog at /admin/instruments.')
bullet('Category filter dropdown — built from existing categories; suggestions: Strings, Brass, Woodwind, Percussion, Keys, Electronic, Vocal, Other.')
bullet('Table columns: Name, Category, Actions.')
bullet('Create / Edit modal: Name (required), Category (text input with datalist).')
spacer()

h1('16. Authors / Contacts')
body('Journalist and contact manager at /admin/authors.')
bullet('Search by name, email, Facebook, or Instagram (real-time).')
bullet('Table: Name, contact channel icons (email, phone, WhatsApp, FB, IG), Notes (truncated), Added date.')
bullet('Create / Edit modal: Name, Email, Phone, WhatsApp, Facebook URL, Instagram URL, Notes textarea.')
bullet('Read-only relations panel in the edit modal shows linked press releases / concerts / tours.')

doc.add_page_break()

# ── 17. Pitch Generator ────────────────────────────────────────────────────────

h1('17. Pitch Generator')
body('AI-assisted email pitch copy tool at /admin/pitch.')
spacer()
h2('Layout')
bullet('Left sidebar: 6 pitch type buttons — Venue, Blog/PR, Playlist, Sync, Festival, Band.')
bullet('Right pane: recipient name input (autocomplete from authors list), custom note textarea, generated pitch preview.')
spacer()
h2('Pitch Generation')
bullet('Automatically fetches: band profile (name, bio_short, genres, comparable artists, stats, booking email), latest 3 upcoming concerts, latest release.')
bullet('Generates template text based on the selected pitch type.')
bullet('"Copy" button copies preview text to clipboard and shows "Copied!" feedback.')
bullet('"Message band" button in Bands table routes here with the band name pre-filled as recipient.')
spacer()
h2('API Calls')
for api in ['GET /api/band-profile', 'GET /api/releases  (latest)', 'GET /api/concerts  (upcoming)', 'GET /api/authors  (autocomplete)']:
    bullet(api)

doc.add_page_break()

# ── 18. Calendar ───────────────────────────────────────────────────────────────

h1('18. Band Calendar')
body('Integrated FullCalendar at /admin/band-calendar aggregating band member personal calendars.')
spacer()
bullet('View toggles: Month (day grid), Week (time grid), List (agenda).')
bullet('Navigation: Previous / Next / Today buttons + current period label.')
bullet('Events sourced from member calendar URLs (iCal) via GET /api/band-profile/calendar/events?start=…&end=….')
bullet('Each member\'s events shown in a distinct colour (8-colour palette).')
bullet('Click an event: popup shows title, start/end, all-day flag, member name/role, description. Dismissible.')
bullet('Members without a calendar_url are silently excluded.')

doc.add_page_break()

# ── 19. Tech Rider ─────────────────────────────────────────────────────────────

h1('19. Tech Rider  (Admin only)')
body('Multi-section technical rider builder at /admin/tech-rider.')
spacer()

h2('Left Sidebar — Template List')
bullet('"New template" button: inline name input + create → POST /api/tech-riders.')
bullet('Each template card: name, "Active" green badge (if active), Edit / Activate / Delete buttons.')
bullet('Activate sets this template as the one linked in the EPK.')
spacer()

h2('Right Pane — Section Editor')
body('Available when a template is open. Top bar: template name (editable), Is Active toggle, Save button.')
spacer()
sections_tech = [
    ('Cover', 'Static band/contact info included at the start of the rider PDF.'),
    ('Stage', 'Drag-and-drop stage plot builder (TechRiderStagePlot component). Add instruments, amps, monitors, FOH position, etc. as draggable icons onto a stage outline.'),
    ('Inputs', 'Channel list table: Channel #, Input type (mic/DI/etc.), Notes. "Suggest inputs" button analyses band member setups to propose a channel count.'),
    ('Monitors', 'Monitor mix table: Monitor #, Sources (which instruments/vocals). "Suggest monitors" button proposes mixes from member setups.'),
    ('Backline', 'Items list with quantity + notes per item (kick drum, snare, hi-hat, toms, bass amp, guitar amps, keyboards, misc). "Suggest backline" auto-populates from member default gear.'),
    ('PA / FOH', 'Form: PA system type, Desired mixing console, Notes.'),
    ('Power', 'Power positions grid + general power notes. "Suggest power positions" button proposes positions from member setups.'),
    ('RF / Wireless', 'Wireless systems list: channel, band, frequency, notes per system. "Suggest" populates from member setups that include wireless gear.'),
]
for name, desc in sections_tech:
    h3(name)
    body(desc)
spacer()

h2('Import Modal')
body('Opens TechRiderImportPanel — select an existing rider template to copy its sections into the current template.')
spacer()

h2('API Calls')
for api in [
    'GET /api/tech-riders',
    'POST /api/tech-riders',
    'GET /api/tech-riders/{id}',
    'PUT /api/tech-riders/{id}',
    'DELETE /api/tech-riders/{id}',
    'POST /api/tech-riders/{id}/activate',
    'GET /api/band-members  (for suggestions)',
    'GET /api/band-member-setups  (for suggestions)',
]:
    bullet(api)

doc.add_page_break()

# ── 20. Newsletter ─────────────────────────────────────────────────────────────

h1('20. Newsletter  (Admin only)')
body('Subscriber list at /admin/newsletter.')
bullet('Header shows total subscriber count.')
bullet('"Export CSV" button downloads all subscribers as a CSV file.')
bullet('Real-time search by email or name.')
bullet('Table: Email (mailto: link), Name, Source badge, Subscribed date, Delete button.')
bullet('Delete: inline Yes / No confirmation (no separate modal).')
bullet('Pagination: Prev / Next / page indicator, shown when total > page size.')
spacer()
h2('API Calls')
for api in ['GET /api/newsletter-subscribers  (paginated)', 'DELETE /api/newsletter-subscribers/{id}']:
    bullet(api)

doc.add_page_break()

# ── 21. Users ──────────────────────────────────────────────────────────────────

h1('21. Users  (Admin only)')
body('User account management at /admin/users.')
spacer()
h2('Table')
bullet('Columns: First name, Last name, Email, Role badge (Admin / Member / Publisher), Created date, Edit / Delete actions.')
spacer()
h2('Create User Modal')
bullet('Band member dropdown (auto-fills first name, last name, email from the linked member).')
bullet('First name, Last name, Email (all required).')
bullet('Password + Password confirmation.')
bullet('Role dropdown: Admin / Member / Publisher.')
spacer()
h2('Edit User Modal')
bullet('Same fields as create; password fields optional (leave blank to keep current password).')
spacer()
h2('API Calls')
for api in ['GET /api/users', 'POST /api/users', 'PUT /api/users/{id}', 'DELETE /api/users/{id}', 'GET /api/band-members  (member dropdown)']:
    bullet(api)

doc.add_page_break()

# ── 22. Setlists ───────────────────────────────────────────────────────────────

h1('22. Setlists  (Admin only)')
body('Setlist builder and song library at /admin/setlists.')
spacer()

h2('Tab 1 — Setlists')
bullet('"New setlist" button: inline name input + create.')
bullet('Setlist list: each entry has Edit, Delete, and "Import from Setlist.fm" buttons.')
bullet('Click a setlist to open the SetlistEditor component.')
h3('SetlistEditor')
bullet('Editable setlist name at the top.')
bullet('"Add songs" button opens song picker from the library.')
bullet('Song list with drag-and-drop reorder.')
bullet('Per-song: title, duration, BPM, key, notes, delete button.')
bullet('Save button sends updated order and song data.')
spacer()

h2('Tab 2 — Song Library')
bullet('"Add song" button opens inline form or modal.')
bullet('Table: Title, Duration (mm:ss), BPM, Key, Notes, Edit / Delete.')
h3('Song Form')
body('Fields: Title (required), Duration in seconds, BPM, Key (e.g. "A minor"), Notes. Save / Cancel.')
spacer()

h2('Setlist.fm Import Modal')
bullet('Paste a Setlist.fm URL.')
bullet('"Fetch" button parses the page client-side.')
bullet('Preview of songs to import shown in a list.')
bullet('"Import" button creates the setlist and all songs, then adds them to the setlist.')
spacer()

h2('API Calls')
for api in [
    'GET /api/setlists', 'POST /api/setlists', 'PUT /api/setlists/{id}', 'DELETE /api/setlists/{id}',
    'GET /api/songs', 'POST /api/songs', 'PUT /api/songs/{id}', 'DELETE /api/songs/{id}',
]:
    bullet(api)

doc.add_page_break()

# ── 23–24. My Profile / My Setups ─────────────────────────────────────────────

h1('23. My Profile  (Member only)')
body('At /admin/my-profile. Uses the same BandMemberForm component as the admin Band Members detail pane, but scoped to the currently authenticated user\'s linked band member.')
bullet('Members can only edit their own record (enforced both in UI and API: PUT /api/band-profile/members/{id} with role:member returns 403 if id ≠ own).')
bullet('Photo upload available.')
bullet('Bio, role, instruments, social links, dates all editable.')
spacer()

h1('24. My Setups  (Member only)')
body('At /admin/my-setups. Uses the same MemberSetupsPanel component as the Band Members Stage Setups tab, scoped to self.')
bullet('Full CRUD on own stage setups.')
bullet('Same section structure as the tech rider builder (inputs, monitors, backline, power, wireless).')

doc.add_page_break()

# ── 25. Shared UI Patterns ─────────────────────────────────────────────────────

h1('25. Shared UI Patterns')
spacer()

h2('Core Components')
table_ncol([
    ['AdminLayout', 'Sidebar nav + main content slot, logout button, user name in footer'],
    ['AdminModal', 'Overlay modal: title prop, open prop, maxWidth prop, @close event'],
    ['ConfirmDialog', 'Delete confirmation overlay: message, loading state, @confirm / @cancel events'],
    ['TableToolbar', 'Search input (v-model) + optional filter slot + "Showing X of Y" counter'],
    ['SortHeader', 'Clickable <th>: label, sortKey, current sort key, direction — emits @sort(key)'],
    ['Pagination', 'Page navigation: page, totalPages, total, perPage, from, to — emits @update:page'],
    ['RichEditor', 'WYSIWYG editor (Quill-based): modelValue / @update:modelValue'],
    ['ImageDropZone', 'Drag-drop or click-to-select image upload, emits File object'],
    ['CareerLevelWidget', 'Career level checklist + level selector buttons (1–4)'],
], ['Component', 'Purpose'])
spacer()

h2('Validation & Error Handling')
bullet('All forms display field-level errors returned from the API.')
bullet('Invalid fields show red error text below the input.')
bullet('Bio tabs show a red indicator dot on the tab label if that sub-tab has a validation error.')
bullet('Submit buttons are disabled and show a spinner while a request is in flight.')
bullet('Successful saves: modal closes (if in modal) or shows "Saved ✓" feedback (if inline).')
bullet('Failed saves: modal stays open with errors visible.')
bullet('Toast notifications: green for success, red for error.')
spacer()

h2('Loading States')
bullet('Tables show a skeleton loader while data is fetching.')
bullet('Modals show a spinner overlay while submitting.')
bullet('Buttons with destructive actions (delete, discard) show a spinner and are disabled during the request.')

doc.add_page_break()

# ── 26. Common Workflows ───────────────────────────────────────────────────────

h1('26. Common Workflows')
body('Key end-to-end flows a QA engineer should cover in test scenarios.')
spacer()

workflows = [
    ('Create a new band member', [
        'Navigate to /admin/band-members.',
        'Click "+" button.',
        'Fill in First name, Last name (required); optionally add Role and Bio.',
        'Click Save → member appears in "Current lineup" sidebar.',
        'Click the new member → select Profile tab → upload photo.',
        'Add instruments from the multi-select.',
        'Add a social link (e.g. Instagram).',
        'Click Save.',
    ]),
    ('Set up the full band profile', [
        'Navigate to /admin/band-profile.',
        'Tab Bio: fill Band name + at least the One-liner bio.',
        'Tab Career: select a level, fill Hometown, Genres.',
        'Tab Social: add at least one link.',
        'Tab Contacts: fill Booking email.',
        'Tab Stats: sync Facebook likes.',
        'Tab EPK: upload tech rider PDF + stage plot image, select featured release.',
        'Save each tab.',
    ]),
    ('Create a concert with poster', [
        'Navigate to /admin/concerts.',
        'Click "Add concert".',
        'Select a venue (must exist — create one first if needed).',
        'Set Date and Start time.',
        'Add bands to the lineup.',
        'Upload a poster image.',
        'Save → concert appears in the Upcoming section.',
    ]),
    ('Publish a blog post with tags', [
        'Navigate to /admin/posts.',
        'Click "Add post".',
        'Enter Title (slug auto-generates).',
        'Write Body in the rich editor.',
        'Select Tags.',
        'Set Published date (today for immediate publish).',
        'Save → post listed under Published.',
    ]),
    ('Add a press article', [
        'Navigate to /admin/press-releases.',
        'Click "Add press release".',
        'Paste the article URL → click "Fetch metadata".',
        'Review auto-filled title / image / description.',
        'Set Published date.',
        'Toggle Featured if article should appear in EPK.',
        'Add tags and relate to concerts or releases.',
        'Save.',
    ]),
    ('Generate and copy a pitch email', [
        'Navigate to /admin/pitch.',
        'Select pitch type (e.g. Venue).',
        'Type or autocomplete recipient name.',
        'Review generated preview (pulls band bio, latest concerts, latest release).',
        'Add a custom note.',
        'Click "Copy" → paste into email client.',
    ]),
    ('Build and activate a tech rider', [
        'Navigate to /admin/tech-rider.',
        'Click "New template", enter a name.',
        'Open the template → fill Cover section.',
        'Go to Stage tab → drag instruments onto stage plot.',
        'Go to Inputs tab → click "Suggest inputs" → review proposed channels.',
        'Fill Monitors, Backline, PA/FOH, Power, RF sections.',
        'Click "Activate" to mark as the active rider for the EPK.',
    ]),
    ('Create and publish an EPK snapshot', [
        'Ensure band profile is complete (bio, social, contacts, career, EPK settings).',
        'Navigate to /admin/band-profile → EPK Settings tab.',
        'Click "Create EPK snapshot" → enter optional reason → Save.',
        'Navigate to /admin (dashboard).',
        'In the EPK Versions widget, click "Publish" on the pending version.',
    ]),
    ('Manage the newsletter list', [
        'Navigate to /admin/newsletter.',
        'Search for a subscriber by email.',
        'Click "Export CSV" to download the full list.',
        'Delete a subscriber using the inline Yes / No confirm.',
    ]),
    ('Import a setlist from Setlist.fm', [
        'Navigate to /admin/setlists.',
        'Click "Import from Setlist.fm" on an existing setlist (or create one first).',
        'Paste the Setlist.fm URL → click "Fetch".',
        'Review the song list preview.',
        'Click "Import" → songs are created and added to the setlist.',
    ]),
]

for wf_title, steps in workflows:
    h2(wf_title)
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}', style='List Number')
    spacer()

# ── Save ───────────────────────────────────────────────────────────────────────

output_path = r'C:\Projects\bandms\admin-panel-description.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
